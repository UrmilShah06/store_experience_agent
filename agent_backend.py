# agent_backend.py
# LangGraph 4-node workflow — footfall, sales, VM/planogram, action plan
# interrupt before process_decisions so store manager can approve/reject actions
#
# Few Thoughts — next iteration
#
# - Add trend comparison for IoT and Sales Performance — right now we only look at current week, would be more
#   useful to show WoW and MoM movement so the store manager sees direction not just snapshot
#
# - Add a store comparison mode — flag when one store's VM compliance is significantly
#   worse than the regional average, useful for area managers not just individual store managers

# - Save previous 3-4 week files of recommended actions, this will be looked up and mapped with 3-4 week sales performance
#   The upcoming week recommendation will consider past weeks recommended actions + sales performance
#   
# - Master Plannogram will be per store 
# 
# - one thing I keep going back to — the action plan right now is static, it doesn't know
#   what was approved last week. need to feed approved_actions history back in as context
#   so the agent stops recommending the same gondola change three weeks in a row



import os
import json
import base64
from pathlib import Path
from typing import TypedDict, Annotated, List, Optional
import operator

import pandas as pd
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver


class StoreExperienceState(TypedDict):
    store_filter: str
    week_filter: str
    api_key: str
    footfall_analysis: str
    sales_analysis: str
    vm_analysis: str
    action_plan: str
    human_decisions: dict
    approved_actions: List[dict]
    rejected_actions: List[dict]
    final_report: str
    session_log: List[str]
    current_node: str
    error: Optional[str]


DATA_DIR = "./data"
_cache = {}  # simple cache — reloads if server restarts, fine for demo


def _fmt_pct(value, decimals=1):
    # wrapping this because I kept writing round(x*100,1) inline everywhere
    # and got it wrong twice — once forgot the *100, once passed already-multiplied value
    # wrapper makes it explicit
    try:
        return round(float(value) * 100, decimals)
    except (TypeError, ValueError):
        return 0.0

# tried using _fmt_pct in the brand_text formatting below to standardise
# but the *100 only applies to the post-fix decimal columns, not all percentages
# IoT columns store already-multiplied values so it would have broken Agent 1
# leaving it here, useful for the sell-through display in Agent 2


def load_data(source):
    if source not in _cache:
        files = {
            "iot": "05_IoT_Store_Behaviour.xlsx",
            "gondola": "06_Gondola_Allocation.xlsx",
            "promo": "07_Promo_Calendar.xlsx",
            "sales": "01_Weekly_Sales_Report.xlsx",
            "sku": "04_SKU_Performance_Report.xlsx",
        }
        path = Path(DATA_DIR) / files[source]
        df = pd.read_excel(path)
        df.columns = df.columns.str.strip()
        _cache[source] = df
    return _cache[source].copy()


def get_latest_week(df):
    return sorted(df['Week'].unique())[-1]


def encode_pdf_page(pdf_path):
    # converts planogram PDF first page to base64 for GPT-4o vision
    # pdf2image needs poppler installed — falls back gracefully if not available
    try:
        from pdf2image import convert_from_path
        import io
        pages = convert_from_path(pdf_path, first_page=1, last_page=1)
        if pages:
            buf = io.BytesIO()
            pages[0].save(buf, format="PNG")
            return base64.b64encode(buf.getvalue()).decode("utf-8")
    except Exception:
        pass
    return ""


def footfall_analysis_node(state):
    # Agent 1 — reads IoT camera data, finds opportunity zones and friction points
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1,
                     openai_api_key=state["api_key"])

    df = load_data("iot")
    store = state.get("store_filter", "all")
    week = state.get("week_filter", "latest")

    if week == "latest":
        target_week = get_latest_week(df)
    else:
        target_week = week

    mask = df['Week'] == target_week
    if store != "all":
        mask &= df['Store'].str.lower() == store.lower()

    filtered = df[mask].copy()

    if filtered.empty:
        state["footfall_analysis"] = f"No IoT data found for store={store}, week={target_week}"
        state["session_log"].append("Agent 1 (Footfall): No data found")
        return state

    # aggregate by zone and category
    
    agg = filtered.groupby(['Zone', 'Category']).agg(
        Total_Visitors=('Visitor_Count', 'sum'),
        Avg_Dwell_Sec=('Avg_Dwell_Time_Sec', 'mean'),
        Avg_Journey=('Product_Journey_Count', 'mean'),
        Avg_Conversion=('Basket_Conversion_Rate_%', 'mean'),
        Avg_Competitor_Int=('Competitor_Interaction_%', 'mean'),
        Female_Pct=('Gender_Female_%', 'mean'),
        Male_Pct=('Gender_Male_%', 'mean'),
        Peak_Visits=('Peak_Hour_Flag', lambda x: (x=='Y').sum()),
    ).reset_index()

    data_text = f"IoT BEHAVIOURAL DATA — {target_week}"
    if store != "all":
        data_text += f" | Store: {store}"
    data_text += f"\nTotal observation records: {len(filtered)}\n\n"

    for _, row in agg.iterrows():
        data_text += (
            f"Zone: {row['Zone']} | Category: {row['Category']}\n"
            f"  Visitors: {int(row['Total_Visitors'])} | "
            f"Avg Dwell: {int(row['Avg_Dwell_Sec'])}s | "
            f"Avg Product Journey: {row['Avg_Journey']:.1f} products touched\n"
            f"  Basket Conversion: {row['Avg_Conversion']:.1f}% | "
            f"Competitor Interaction: {row['Avg_Competitor_Int']:.1f}%\n"
            f"  Gender Split: F={row['Female_Pct']:.0f}% M={row['Male_Pct']:.0f}%\n\n"
        )
    # not passing Peak_Visits to the LLM — tried it and the agent kept fixating on
    # peak hour staffing recommendations instead of VM actions, which is out of scope

    response = llm.invoke([
        SystemMessage(content="""You are a Retail Footfall Intelligence Analyst specialising in 
IoT camera data interpretation. You have expertise in customer journey analytics, 
behavioural patterns, and conversion optimisation in large format retail stores.
Your analysis should identify: high-opportunity zones (high footfall, low conversion),
friction points (high dwell, low purchase), demographic mismatches with product range,
and peak hour patterns that should influence VM decisions."""),
        HumanMessage(content=f"""Analyse this IoT behavioural data and identify key insights:

{data_text}

Provide:
1. TOP 3 OPPORTUNITY ZONES: High visitor count but low basket conversion (untapped potential)
2. FRICTION POINTS: High dwell time but low conversion (customer interested but not buying)
3. DEMOGRAPHIC INSIGHTS: Where gender/age mix suggests VM or range adjustments
4. PEAK HOUR PATTERNS: Which zones need extra VM attention during peak hours
5. COMPETITOR THREAT ZONES: Where competitor interaction is high (risk of losing customers)

Be specific with zone names, category names, and percentages from the data.""")
    ])

    state["footfall_analysis"] = response.content
    state["session_log"].append(f"Agent 1 (Footfall Analysis): Completed — {target_week}")
    state["current_node"] = "footfall_analysis"
    return state


def sales_analysis_node(state):
    # Agent 2 — correlates sales with footfall to find VM opportunities
    # the key insight here: a brand with high sell-through but low sales revenue
    # is likely undershelved — limited gondola space means limited stock displayed,
    # which caps sales even when customer demand is strong
    # it's not an inventory replenishment problem, it's a space allocation problem —
    # the fix is more facings or a better shelf position, not a bigger PO
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1,
                     openai_api_key=state["api_key"])

    sales_df = load_data("sales")
    sku_df = load_data("sku")
    store = state.get("store_filter", "all")
    week = state.get("week_filter", "latest")

    if week == "latest":
        target_week = get_latest_week(sales_df)  
    else:
        target_week = week

    mask = sales_df['Week'] == target_week
    if store != "all":
        mask &= sales_df['Store'].str.lower() == store.lower()

    sales = sales_df[mask].copy()

    cat_perf = []
    for _, row in sales.iterrows():
        try:
            actual = float(row['Actual_INR_L'])
            target = float(row['Target_INR_L'])
            ach = round(actual/target*100, 1) if target > 0 else 0
            lw = float(row['LW_INR_L'])
            wow = round((actual/lw-1)*100, 1) if lw > 0 else 0
            cat_perf.append(f"{row['Category']} @ {row['Store']}: "
                           f"INR {actual}L, Ach {ach}%, WoW {wow:+.1f}%")
        except Exception:
            pass  # skip malformed rows

    sku_mask = sku_df['Week'] == target_week
    if store != "all":
        sku_mask &= sku_df['Store'].str.lower() == store.lower()

    sku_filtered = sku_df[sku_mask].copy()
    brand_agg = sku_filtered.groupby(['Category', 'Brand']).agg(
        Units=('Units_Sold', 'sum'),
        Avg_ST=('Sell_Through_%', 'mean'),
        Low_Stock=('Days_Cover', lambda x: (pd.to_numeric(x, errors='coerce') < 14).sum())
    ).reset_index().sort_values('Avg_ST', ascending=False)

    brand_text = ""
    for _, row in brand_agg.head(20).iterrows():
        brand_text += (f"{row['Brand']} ({row['Category']}): "
                      f"{int(row['Units'])} units, ST {row['Avg_ST']*100:.1f}%, "
                      f"Low stock SKUs: {int(row['Low_Stock'])}\n")

    # tried sorting by Low_Stock descending instead of Avg_ST to surface reorder alerts first
    # but then the LLM led with stock replenishment recommendations every time
    # and buried the VM opportunity brands — ST sort gives better action plan output
    # brand_agg_by_stock = brand_agg.sort_values('Low_Stock', ascending=False)

    # pass first 500 chars of footfall output as context — enough for signal without overwhelming
    footfall_context = state.get("footfall_analysis", "")[:500]

    response = llm.invoke([
        SystemMessage(content="""You are a Senior Retail Sales Performance Analyst.
You specialise in correlating customer behaviour data with sales outcomes to identify
where Visual Merchandising changes would have the highest revenue impact.
Focus on: brands performing above expectations despite poor placement,
categories with strong footfall but weak sales (VM opportunity),
and brands with high sell-through that need more gondola space."""),
        HumanMessage(content=f"""Correlate sales performance with IoT footfall patterns.

SALES DATA ({target_week}):
{chr(10).join(cat_perf[:15])}

BRAND SELL-THROUGH:
{brand_text}

FOOTFALL CONTEXT (from Agent 1):
{footfall_context}

Identify:
1. VM OPPORTUNITY BRANDS: High sell-through but likely poor placement (needs more space/better position)
2. CATEGORY MISMATCHES: High footfall zone but underperforming category (wrong VM for the audience)
3. QUICK WIN GONDOLA CHANGES: Specific brand repositioning that would lift sales
4. STOCK ALERT BRANDS: High sell-through + low days cover (risk of losing sales due to stockout)

Use specific brand names, category names, and numbers from the data.""")
    ])

    state["sales_analysis"] = response.content
    state["session_log"].append(f"Agent 2 (Sales Analysis): Completed — {target_week}")
    state["current_node"] = "sales_analysis"
    return state


def vm_analysis_node(state):
    # Agent 3 — multimodal node, reads gondola Excel + planogram PDF via GPT-4o Vision
    
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1,
                     openai_api_key=state["api_key"])

    gondola_df = load_data("gondola")
    store = state.get("store_filter", "all")

    if store != "all":
        gondola = gondola_df[gondola_df['Store'].str.lower() == store.lower()].copy()
    else:
        gondola = gondola_df.copy()

    compliance_gaps = gondola[gondola['Compliance_Gap'] == 'YES']
    compliance_text = f"Gondola Allocation Analysis:\n"
    compliance_text += f"Total fixtures: {len(gondola)} | Compliance gaps: {len(compliance_gaps)}\n\n"

    for _, row in compliance_gaps.head(15).iterrows():
        compliance_text += (
            f"GAP: {row['Brand']} in {row['Zone']} ({row['Store']})\n"
            f"  Current: {row['Shelf_Level']} | Required: {row['Planogram_Specified_Level']}\n"
            f"  VM Score: {row['VM_Compliance_Score_%']}%\n\n"
        )

    planogram_insight = ""
    pdf_path = str(Path(DATA_DIR) / "08_Store_Planogram.pdf")
    img_b64 = encode_pdf_page(pdf_path)

    if img_b64:
        try:
            vision_response = llm.invoke([
                SystemMessage(content="You are a VM planogram expert. Analyse retail store floor plans."),
                HumanMessage(content=[
                    {"type": "image_url",
                     "image_url": {"url": f"data:image/png;base64,{img_b64}", "detail": "high"}},
                    {"type": "text",
                     "text": """Analyse this retail store planogram floor plan.
Extract:
1. Which categories are assigned to which zones (Zone A, B, C, D, E)
2. Which brands are specified for premium positions (eye level, front bays)
3. Any special fixtures mentioned (shop-in-shop, brand walls, seated display)
4. VM compliance rules stated in the guidelines table
Be specific about zone names and brand names visible in the document."""}
                ])
            ])
            planogram_insight = f"\nPLANOGRAM ANALYSIS (GPT-4o Vision):\n{vision_response.content}"
        except Exception as e:
            planogram_insight = f"\nPlanogram vision analysis unavailable: {e}"
    else:
        # vision not available — fall back to hardcoded guidelines
        # TODO: load these from a config file storewise, instead of hardcoding
        # each store genuinely has a different planogram — Koramangala has a different
        # zone layout to Whitefield, hardcoding one set of rules is a known gap here
        planogram_insight = """
PLANOGRAM GUIDELINES (Standard):
- Zone A (Entrance): Ladies Western front bay — AND/W brands eye level priority
- Zone B Left: Men Formal — Van Heusen/Arrow eye level mandatory
- Zone C Centre: Ladies Ethnic — Biba shop-in-shop, minimum 4 linear feet
- Zone D Right: Ladies/Men Footwear — Premium brands at zone entrance
- Zone E Back: Sportswear — Nike/Adidas separate brand walls"""

    footfall_context = state.get("footfall_analysis", "")[:400]
    sales_context = state.get("sales_analysis", "")[:400]

    response = llm.invoke([
        SystemMessage(content="""You are a Senior Visual Merchandising and Planogram Compliance Analyst.
You review gondola allocation data against planogram guidelines to produce
prioritised VM action recommendations. Focus on compliance gaps that are
costing the most revenue, and quick wins that require minimal disruption."""),
        HumanMessage(content=f"""Review gondola compliance and recommend VM actions.

{compliance_text}

{planogram_insight}

FOOTFALL CONTEXT: {footfall_context}
SALES CONTEXT: {sales_context}

Produce:
1. CRITICAL COMPLIANCE GAPS: Brands in wrong shelf position that are hurting conversion
2. QUICK WIN VM CHANGES: Gondola adjustments that can be done this week
3. FACE COUNT RECOMMENDATIONS: Brands that deserve more/less gondola space
4. ZONE OPTIMISATION: Any zone-level changes to better match customer flow
5. VM CALENDAR ALIGNMENT: VM actions to prepare for upcoming promotions

Prioritise by revenue impact. Be specific with store names, brand names, zone names.""")
    ])

    state["vm_analysis"] = response.content
    state["session_log"].append("Agent 3 (VM Analysis): Completed — Planogram reviewed")
    state["current_node"] = "vm_analysis"
    return state


def action_plan_node(state):
    # Agent 4 — pulls everything together into a ranked 7-action monthly plan
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1,
                     openai_api_key=state["api_key"])

    promo_df = load_data("promo")
    store = state.get("store_filter", "all")
    week = state.get("week_filter", "latest")

    if week == "latest":
        target_week = get_latest_week(promo_df)
    else:
        target_week = week

    weeks_sorted = sorted(promo_df['Week'].unique())
    try:
        curr_idx = weeks_sorted.index(target_week)
        upcoming_weeks = weeks_sorted[curr_idx:curr_idx+4]
    except ValueError:
        upcoming_weeks = weeks_sorted[-4:]

    promo_mask = promo_df['Week'].isin(upcoming_weeks)
    if store != "all":
        promo_mask &= (promo_df['Store'].str.lower() == store.lower()) | \
                      (promo_df['Store_Scope'].isin(['Chain', 'Regional']))

    upcoming_promos = promo_df[promo_mask].copy()
    promo_text = "UPCOMING PROMOTIONS (next 4 weeks):\n"
    for _, row in upcoming_promos.head(10).iterrows():
        promo_text += (f"  {row['Week']}: {row['Promo_Name']} — {row['Category']} | "
                      f"Type: {row['Offer_Type']} | Priority: {row['Priority_Level']} | "
                      f"VM Support: {row['VM_Support_Required']}\n")

    response = llm.invoke([
        SystemMessage(content="""You are a Senior Retail Store Operations Consultant.
You synthesise IoT behavioural data, sales performance, VM compliance analysis,
and the promotional calendar into a concrete monthly store action plan.
Each action must have: priority, owner (Store Manager/Category Manager/VM Team),
effort (hours), expected impact, and specific details.
Actions must be executable — not generic advice."""),
        HumanMessage(content=f"""Synthesise all analyses into a monthly store action plan.

FOOTFALL ANALYSIS SUMMARY:
{state.get('footfall_analysis', '')[:600]}

SALES ANALYSIS SUMMARY:
{state.get('sales_analysis', '')[:600]}

VM ANALYSIS SUMMARY:
{state.get('vm_analysis', '')[:600]}

{promo_text}

Generate a MONTHLY STORE ACTION PLAN with exactly this structure:

**EXECUTIVE SUMMARY**
2-3 sentences on the biggest opportunity this month.

**PRIORITY ACTIONS** (numbered 1-7, sorted by revenue impact):
For each action:
- Action: [specific gondola/VM/promo action]
- Category/Brand: [specific names]
- Zone/Fixture: [specific location]
- Owner: [Store Manager / Category Manager / VM Team]
- Effort: [estimated hours]
- Expected Impact: [projected conversion uplift or revenue]
- Rationale: [1 line — which data signal drives this]

**PROMO READINESS**
VM actions needed before each upcoming promotion.

**METRICS TO TRACK**
3 KPIs to measure success of these actions next week.""")
    ])

    state["action_plan"] = response.content
    state["session_log"].append("Agent 4 (Action Plan): Draft completed")
    state["current_node"] = "action_plan"
    return state


def human_approval_node(state):
    # pause point — LangGraph interrupt mechanism kicks in here
    # Streamlit checks current_node == "awaiting_approval" to show the review UI
    state["current_node"] = "awaiting_approval"
    state["session_log"].append("Human approval checkpoint reached")
    return state


def process_decisions_node(state):
    # processes store manager approvals and logs via MCP tools
    from tools import log_vm_action, get_store_layout

    decisions = state.get("human_decisions", {})
    approved = []
    rejected = []

    for action_id, decision in decisions.items():
        if decision["status"] in ["approve", "modified"]:
            action_record = {
                "action_id": action_id,
                "description": decision.get("description", ""),
                "quantity": decision.get("quantity", ""),
                "status": decision["status"],
                "approved_by": "Store Manager",
            }
            result = log_vm_action(
                store=state.get("store_filter", "all"),
                action_id=action_id,
                description=decision.get("description", ""),
                status=decision["status"]
            )
            action_record["tool_result"] = result
            approved.append(action_record)
        else:
            rejected.append({
                "action_id": action_id,
                "reason": decision.get("reason", "Rejected by manager"),
            })

    state["approved_actions"] = approved
    state["rejected_actions"] = rejected
    state["session_log"].append(
        f"Decisions processed: {len(approved)} approved, {len(rejected)} rejected"
    )
    state["current_node"] = "decisions_processed"
    return state


def generate_report_node(state):
    llm = ChatOpenAI(model="gpt-4o", temperature=0.1,
                     openai_api_key=state["api_key"])

    approved = state.get("approved_actions", [])
    rejected = state.get("rejected_actions", [])
    # keeping this node simple — it just summarises what happened
    # tried making it pull IoT + sales data again for a richer summary
    # but that's another 2 LLM calls and the manager just wants to see what got approved

    approved_text = "\n".join([f"- {a['action_id']}: {a['description']} ({a['status']})"
                               for a in approved]) or "None"
    rejected_text = "\n".join([f"- {r['action_id']}: {r.get('reason','')}"
                               for r in rejected]) or "None"

    response = llm.invoke([
        SystemMessage(content="You are a retail operations report writer."),
        HumanMessage(content=f"""Generate a concise session completion report.

APPROVED ACTIONS ({len(approved)}):
{approved_text}

REJECTED ACTIONS ({len(rejected)}):
{rejected_text}

ORIGINAL ACTION PLAN:
{state.get('action_plan', '')[:400]}

Write a 1-page session summary:
- What was analysed
- Key findings from IoT, sales, and VM data
- Actions approved and their expected combined impact
- Actions rejected and why this matters
- What to measure next week to track progress""")
    ])

    state["final_report"] = response.content
    state["session_log"].append("Final report generated")
    state["current_node"] = "complete"
    return state


def build_workflow():
    # sequential: footfall -> sales -> vm -> action_plan -> approval -> decisions -> report
    workflow = StateGraph(StoreExperienceState)

    workflow.add_node("footfall_analysis_node", footfall_analysis_node)
    workflow.add_node("sales_analysis_node", sales_analysis_node)
    workflow.add_node("vm_analysis_node", vm_analysis_node)
    workflow.add_node("action_plan_node", action_plan_node)
    workflow.add_node("human_approval_node", human_approval_node)
    workflow.add_node("process_decisions_node", process_decisions_node)
    workflow.add_node("generate_report_node", generate_report_node)

    workflow.set_entry_point("footfall_analysis_node")
    workflow.add_edge("footfall_analysis_node", "sales_analysis_node")
    workflow.add_edge("sales_analysis_node", "vm_analysis_node")
    workflow.add_edge("vm_analysis_node", "action_plan_node")
    workflow.add_edge("action_plan_node", "human_approval_node")
    workflow.add_edge("human_approval_node", "process_decisions_node")
    workflow.add_edge("process_decisions_node", "generate_report_node")
    workflow.add_edge("generate_report_node", END)

    memory = MemorySaver()
    return workflow.compile(
        checkpointer=memory,
        interrupt_before=["process_decisions_node"]
    )


def get_initial_state(store, week, api_key):
    return StoreExperienceState(
        store_filter=store,
        week_filter=week,
        api_key=api_key,
        footfall_analysis="",
        sales_analysis="",
        vm_analysis="",
        action_plan="",
        human_decisions={},
        approved_actions=[],
        rejected_actions=[],
        final_report="",
        session_log=[],
        current_node="start",
        error=None,
    )


def generate_word_document(store, week, agent_outputs, approved_actions, rejected_actions, final_report):
    # generates session report as Word doc
    # had to handle the case where approved_actions is empty list — table was erroring on zero rows
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading("STORE EXPERIENCE OPTIMISATION — SESSION REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        subtitle = doc.add_paragraph(f"Store: {store}  |  Week: {week}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(11)
        subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        doc.add_paragraph("─" * 80)

        sections = [
            ("AGENT 1 — FOOTFALL INTELLIGENCE ANALYSIS", agent_outputs.get("footfall_analysis", "")),
            ("AGENT 2 — SALES PERFORMANCE ANALYSIS",     agent_outputs.get("sales_analysis", "")),
            ("AGENT 3 — VM & PLANOGRAM ANALYSIS",        agent_outputs.get("vm_analysis", "")),
            ("AGENT 4 — MONTHLY ACTION PLAN",            agent_outputs.get("action_plan", "")),
            ("FINAL SESSION REPORT",                     final_report),
        ]

        for section_title, content in sections:
            if not content:
                continue
            h = doc.add_heading(section_title, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith('**') and line.endswith('**'):
                    ph = doc.add_heading(line.strip('*'), level=2)
                    ph.runs[0].font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
                elif line.startswith('- ') or line.startswith('• '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)

        # approved actions table
        if approved_actions:
            h = doc.add_heading("APPROVED VM ACTIONS", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, text in enumerate(["Action ID", "Description", "Status"]):
                hdr[i].text = text
                hdr[i].paragraphs[0].runs[0].font.bold = True
            for action in approved_actions:
                row = table.add_row().cells
                row[0].text = action.get("action_id", "")
                row[1].text = action.get("description", "")
                row[2].text = action.get("status", "")

        # rejected actions table
        if rejected_actions:
            h = doc.add_heading("REJECTED ACTIONS", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            for i, text in enumerate(["Action ID", "Reason"]):
                hdr[i].text = text
                hdr[i].paragraphs[0].runs[0].font.bold = True
            for action in rejected_actions:
                row = table.add_row().cells
                row[0].text = action.get("action_id", "")
                row[1].text = action.get("reason", "Rejected by manager")

        out_path = "./data/Store_Experience_Session_Report.docx"
        doc.save(out_path)
        return out_path

    except ImportError:
        out_path = "./data/Store_Experience_Session_Report.txt"
        with open(out_path, 'w') as f:
            f.write(final_report)
        return out_path
